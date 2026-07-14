"""统一文档解析适配器。

解析器只负责把文件转换为文本、表格和来源定位，不做业务判断。
支持 PDF、Word、文本、Markdown、CSV 和 Excel；缺少可选依赖时给出明确提示。
"""

from __future__ import annotations

import csv
import io
from dataclasses import dataclass, field
from pathlib import Path
from typing import Any


@dataclass(slots=True)
class ParsedDocument:
    """一个文件解析后的规范化表示。"""

    name: str
    text: str
    tables: list[list[dict[str, Any]]] = field(default_factory=list)
    metadata: dict[str, Any] = field(default_factory=dict)

    def evidence_for(self, needle: str, context: int = 80) -> tuple[str, str]:
        """在文本中定位关键词并返回行号与截断片段。"""

        if not needle:
            return "全文", self.text[: context * 2]
        lowered = self.text.lower()
        index = lowered.find(needle.lower())
        if index < 0:
            return "全文", self.text[: context * 2]
        line_number = self.text[:index].count("\n") + 1
        start = max(0, index - context)
        end = min(len(self.text), index + len(needle) + context)
        return f"第 {line_number} 行", self.text[start:end].replace("\n", " ")


def parse_document(path: str | Path) -> ParsedDocument:
    """按扩展名解析文件并返回统一文档对象。"""

    file_path = Path(path)
    if not file_path.is_file():
        raise FileNotFoundError(f"文件不存在: {file_path}")
    suffix = file_path.suffix.lower()
    if suffix in {".txt", ".md"}:
        return _parse_text(file_path)
    if suffix == ".pdf":
        return _parse_pdf(file_path)
    if suffix == ".docx":
        return _parse_docx(file_path)
    if suffix == ".csv":
        return _parse_csv(file_path)
    if suffix in {".xlsx", ".xlsm"}:
        return _parse_excel(file_path)
    raise ValueError(f"暂不支持 {suffix or '无扩展名'} 文件，请使用 PDF/DOCX/TXT/MD/CSV/XLSX")


def _read_text_bytes(content: bytes) -> str:
    """按常见中文编码顺序解码文本。"""

    for encoding in ("utf-8-sig", "utf-8", "gb18030", "gbk"):
        try:
            return content.decode(encoding)
        except UnicodeDecodeError:
            continue
    return content.decode("utf-8", errors="replace")


def _parse_text(path: Path) -> ParsedDocument:
    """解析纯文本或 Markdown。"""

    text = _read_text_bytes(path.read_bytes())
    return ParsedDocument(
        path.name, text, metadata={"type": "text", "line_count": len(text.splitlines())}
    )


def _parse_pdf(path: Path) -> ParsedDocument:
    """逐页解析 PDF，并在文本中保留页码边界。"""

    try:
        from pypdf import PdfReader
    except ImportError as exc:
        raise ImportError("解析 PDF 需要安装 pypdf") from exc
    reader = PdfReader(str(path))
    pages: list[str] = []
    for page_number, page in enumerate(reader.pages, start=1):
        pages.append(f"\n--- 第 {page_number} 页 ---\n{page.extract_text() or ''}")
    return ParsedDocument(
        path.name,
        "".join(pages).strip(),
        metadata={"type": "pdf", "page_count": len(reader.pages)},
    )


def _parse_docx(path: Path) -> ParsedDocument:
    """解析 Word 段落和表格。"""

    try:
        from docx import Document
    except ImportError as exc:
        raise ImportError("解析 DOCX 需要安装 python-docx") from exc
    document = Document(str(path))
    paragraphs = [paragraph.text for paragraph in document.paragraphs if paragraph.text.strip()]
    tables: list[list[dict[str, Any]]] = []
    for table in document.tables:
        rows = [[cell.text.strip() for cell in row.cells] for row in table.rows]
        if not rows:
            continue
        headers = _unique_headers(rows[0])
        tables.append([dict(zip(headers, row, strict=False)) for row in rows[1:]])
    table_text = "\n".join(_records_to_text(table) for table in tables)
    text = "\n".join(paragraphs + ([table_text] if table_text else []))
    return ParsedDocument(
        path.name,
        text,
        tables=tables,
        metadata={"type": "docx", "paragraph_count": len(paragraphs)},
    )


def _parse_csv(path: Path) -> ParsedDocument:
    """解析 CSV，保留每一行字段名和值。"""

    text = _read_text_bytes(path.read_bytes())
    sample = text[:4096]
    try:
        dialect = csv.Sniffer().sniff(sample, delimiters=",\t;|")
    except csv.Error:
        dialect = csv.excel
    reader = csv.DictReader(io.StringIO(text), dialect=dialect)
    records = [{str(key).strip(): value for key, value in row.items()} for row in reader]
    return ParsedDocument(
        path.name,
        _records_to_text(records),
        tables=[records],
        metadata={"type": "csv", "row_count": len(records)},
    )


def _parse_excel(path: Path) -> ParsedDocument:
    """解析 Excel 的所有非空工作表。"""

    try:
        from openpyxl import load_workbook
    except ImportError as exc:
        raise ImportError("解析 Excel 需要安装 openpyxl") from exc
    workbook = load_workbook(path, data_only=True, read_only=True)
    tables: list[list[dict[str, Any]]] = []
    sheets: list[str] = []
    text_parts: list[str] = []
    for worksheet in workbook.worksheets:
        rows = list(worksheet.iter_rows(values_only=True))
        rows = [row for row in rows if any(value not in (None, "") for value in row)]
        if not rows:
            continue
        headers = _unique_headers([str(value or "").strip() for value in rows[0]])
        records = [dict(zip(headers, row, strict=False)) for row in rows[1:]]
        tables.append(records)
        sheets.append(worksheet.title)
        text_parts.append(f"工作表: {worksheet.title}\n{_records_to_text(records)}")
    workbook.close()
    return ParsedDocument(
        path.name,
        "\n\n".join(text_parts),
        tables=tables,
        metadata={"type": "excel", "sheets": sheets, "row_count": sum(map(len, tables))},
    )


def _unique_headers(headers: list[str]) -> list[str]:
    """把空表头和重复表头转换为稳定唯一名称。"""

    seen: dict[str, int] = {}
    result: list[str] = []
    for index, header in enumerate(headers, start=1):
        base = header or f"列{index}"
        seen[base] = seen.get(base, 0) + 1
        result.append(base if seen[base] == 1 else f"{base}_{seen[base]}")
    return result


def _records_to_text(records: list[dict[str, Any]]) -> str:
    """把表格记录转换为适合规则和模型读取的逐行文本。"""

    lines: list[str] = []
    for index, record in enumerate(records, start=2):
        values = "；".join(f"{key}={value}" for key, value in record.items())
        lines.append(f"第{index}行：{values}")
    return "\n".join(lines)
